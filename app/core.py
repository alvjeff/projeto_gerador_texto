import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from jinja2 import Environment, FileSystemLoader


def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except AttributeError:
        base_path = Path(__file__).resolve().parent

    return Path(base_path) / relative_path


def gerar_documento(dados, nome_arquivo):
    # carregar template
    template_dir = resource_path("templates")
    env = Environment(loader=FileSystemLoader(template_dir))
    template = env.get_template("template.txt")
    texto_final = template.render(dados)

    # pasta de saída (onde o usuário estiver)
    output_dir = Path.cwd() / "output"
    output_dir.mkdir(exist_ok=True)

    # cria documento Word
    doc = Document()

    titulo = doc.add_paragraph()
    run = titulo.add_run("DECLARAÇÃO")
    run.bold = True
    run.font.size = Pt(14)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    for linha in texto_final.split("\n"):
        p = doc.add_paragraph(linha)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(12)

    caminho = output_dir / f"{nome_arquivo}.docx"
    doc.save(caminho)

    return caminho
