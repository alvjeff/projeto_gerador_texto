import os
import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from jinja2 import Environment, FileSystemLoader
from pathlib import Path


# FUNCAO PARA LOCALIZAR RECURSOS
def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS #executavel
    except AttributeError:
        base_path = os.path.abspath(".") #python normal

    return os.path.join(base_path, relative_path)

#coleta dos dados
dados = {
    "nome": input("Nome: "),
    "funcao": input("Função: "),
    "periodo": input("Período: "),
    "unidade": input("Unidade: ")
}

# nome do arquivo
nome_arquivo = input("Nome do arquivo (sem extensão): ").strip()
if not nome_arquivo:
    nome_arquivo = "declaracao"

""""
# carrega o template
env = Environment(loader=FileSystemLoader("."))
template = env.get_template("template.txt")
texto_final = template.render(dados)
"""

#carrega o template (correto para executavel)
template_dir = resource_path("") #pasta base
env = Environment(loader=FileSystemLoader(template_dir))
template = env.get_template("template.txt")
texto_final = template.render(dados)


""""
# gera o texto final
texto_final = template.render(dados)
"""

#pasta de saida (fora do executavel)
output_dir = Path.cwd() / "output"
output_dir.mkdir(exist_ok=True)

"""
# caminhos
BASE_DIR = Path(__file__).resolve().parent
output_dir = BASE_DIR / "output"
output_dir.mkdir(exist_ok=True)
"""

"""
#garante pasta de saída
output_dir = Path("output")
output_dir.mkdir(exist_ok=True)
"""

# cria documento Word
doc = Document()

# TITULO
titulo = doc.add_paragraph()
run_titulo = titulo.add_run("DECLARAÇÃO")
run_titulo.bold = True
run_titulo.font.size = Pt(14)
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph() #linha em branco


#CORPO DO TEXTO
for linha in texto_final.split("\n"):
    p = doc.add_paragraph(linha)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(12)

# salvar
caminho = output_dir / f"{nome_arquivo}.docx"
doc.save(caminho)

print(f"\nDocumento Word gerado com sucesso em: {caminho}")

#salva o arquivo (por enquanto txt)
"""caminho = output_dir / f"{nome_arquivo}.txt"
with open(caminho, "w", encoding="utf-8") as f:
    f.write(texto_final)

print("\nDocumento gerado com sucesso em: {caminho}")"""


"""
#cria word
doc = Document()
for linha in texto_final.split("\n"):
    doc.add_paragraph(linha)

caminho = output_dir / f"{nome_arquivo}.docx"
doc.save(caminho)

print("f\nDocumento Word gerado com sucesso em: {caminho}")
"""